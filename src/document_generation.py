
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.lib import colors
from docx import Document
from .resume_template import generate_professional_template

def generate_pdf(resume_text, file_path):
    """
    Generates a professional, ATS-friendly PDF resume.
    """
    doc = SimpleDocTemplate(
        file_path,
        pagesize=letter,
        rightMargin=72,
        leftMargin=72,
        topMargin=72,
        bottomMargin=72
    )

    styles = getSampleStyleSheet()

    # Custom Styles
    title_style = ParagraphStyle(
        'TitleStyle',
        parent=styles['Heading1'],
        fontSize=18,
        textColor=colors.HexColor("#2C3E50"),
        alignment=1, # Center
        spaceAfter=10
    )

    contact_style = ParagraphStyle(
        'ContactStyle',
        parent=styles['Normal'],
        fontSize=10,
        alignment=1,
        spaceAfter=20
    )

    heading_style = ParagraphStyle(
        'HeadingStyle',
        parent=styles['Heading2'],
        fontSize=14,
        textColor=colors.HexColor("#2C3E50"),
        borderPadding=(0, 0, 2, 0),
        borderWidth=0,
        spaceBefore=12,
        spaceAfter=6
    )

    normal_style = styles['Normal']
    normal_style.fontSize = 11
    normal_style.leading = 14

    sections = generate_professional_template(resume_text)
    story = []

    # Header
    story.append(Paragraph(sections["name"], title_style))
    story.append(Paragraph(sections["contact"], contact_style))

    # Summary
    if sections["summary"]:
        story.append(Paragraph("Summary", heading_style))
        story.append(Paragraph(sections["summary"], normal_style))

    # Skills
    if sections["skills"]:
        story.append(Paragraph("Skills", heading_style))
        # Use columns for skills to save space? No, ATS prefers simple lists.
        # But we can do a comma-separated paragraph or bullets.
        # Bullet list:
        skill_items = [ListItem(Paragraph(skill, normal_style)) for skill in sections["skills"]]
        if skill_items:
            story.append(ListFlowable(skill_items, bulletType='bullet'))

    # Experience
    if sections["experience"]:
        story.append(Paragraph("Experience", heading_style))
        for item in sections["experience"]:
            # Check if item has multiple lines (e.g., bullets)
            lines = item.split('\n')
            for i, line in enumerate(lines):
                if line.strip().startswith('-') or line.strip().startswith('•'):
                    # Bullet point
                    clean_line = line.strip('- •')
                    story.append(ListFlowable([ListItem(Paragraph(clean_line, normal_style))], bulletType='bullet', leftIndent=20))
                else:
                    story.append(Paragraph(line, normal_style))
            story.append(Spacer(1, 0.1 * inch))

    # Education
    if sections["education"]:
        story.append(Paragraph("Education", heading_style))
        for item in sections["education"]:
            story.append(Paragraph(item.replace('\n', '<br/>'), normal_style))
            story.append(Spacer(1, 0.1 * inch))

    doc.build(story)

def generate_docx(resume_text, file_path):
    """
    Generates a professional, ATS-friendly DOCX resume.
    """
    document = Document()
    sections = generate_professional_template(resume_text)

    # Header
    document.add_heading(sections["name"], 0)
    p = document.add_paragraph()
    p.alignment = 1 # Center
    p.add_run(sections["contact"])

    # Summary
    if sections["summary"]:
        document.add_heading('Summary', level=1)
        document.add_paragraph(sections["summary"])

    # Skills
    if sections["skills"]:
        document.add_heading('Skills', level=1)
        for skill in sections["skills"]:
            document.add_paragraph(skill, style='List Bullet')

    # Experience
    if sections["experience"]:
        document.add_heading('Experience', level=1)
        for item in sections["experience"]:
            lines = item.split('\n')
            for line in lines:
                if line.strip().startswith('-') or line.strip().startswith('•'):
                    document.add_paragraph(line.strip('- •'), style='List Bullet')
                else:
                    document.add_paragraph(line)

    # Education
    if sections["education"]:
        document.add_heading('Education', level=1)
        for item in sections["education"]:
            document.add_paragraph(item)

    document.save(file_path)
